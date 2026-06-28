from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs/reports/2026-06-16-marketplace-progress-report.md"
OUTPUT = ROOT / "docs/reports/2026-06-16-marketplace-progress-report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(10)

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.left_indent = Inches(0.18)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    widths = [2600, 6760] if len(rows[0]) == 2 else [9360 // len(rows[0])] * len(rows[0])

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[c_idx])
            text = value.strip()
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    if r_idx == 0:
                        run.font.bold = True
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")

    doc.add_paragraph()


def add_code_block(doc, lines):
    for line in lines:
        paragraph = doc.add_paragraph(style="Code Block")
        run = paragraph.add_run(line.rstrip() or " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("1F2937")


def export():
    doc = Document()
    style_document(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            add_code_block(doc, code_lines)
            code_lines = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            parts = [part.strip() for part in line.strip("|").split("|")]
            if all(set(part) <= {"-", " "} for part in parts):
                continue
            table_rows.append(parts)
            continue

        flush_table()

        if not line:
            continue

        if line.startswith("# "):
            paragraph = doc.add_paragraph(line[2:].strip(), style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(line[2:].strip())
        elif line[0].isdigit() and ". " in line[:4]:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.add_run(line.split(". ", 1)[1].strip())
        else:
            doc.add_paragraph(line)

    flush_code()
    flush_table()

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "码上好 AI 开发交易平台阶段工作报告"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("6B7280")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    export()
